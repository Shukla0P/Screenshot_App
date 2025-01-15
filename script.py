import pyautogui
import tkinter as tk
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
from pynput import keyboard
import threading
import os
import shutil
from datetime import datetime

# Global Variables
temp_folder = "temp"
word_file_path = "screenshots.docx"

# Ensure temp folder exists
if not os.path.exists(temp_folder):
    os.makedirs(temp_folder)

# Clear Temp Folder
def clear_temp_folder():
    for file in os.listdir(temp_folder):
        file_path = os.path.join(temp_folder, file)
        if os.path.isfile(file_path) or os.path.islink(file_path):
            os.unlink(file_path)
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)

# Capture Screenshot
def capture_screenshot():
    clear_temp_folder()  # Clear temp folder before capturing a new screenshot
    timestamp = datetime.now().strftime('%Y_%m_%d_%H_%M_%S')
    screenshot_path = os.path.join(temp_folder, f"screenshot_{timestamp}.png")
    screenshot = pyautogui.screenshot()
    screenshot.save(screenshot_path)
    open_crop_gui(screenshot_path)

# Capture Full-Screen Screenshot
def capture_fullscreen_screenshot():
    clear_temp_folder()  # Clear temp folder before capturing a new screenshot
    timestamp = datetime.now().strftime('%Y_%m_%d_%H_%M_%S')
    screenshot_path = os.path.join(temp_folder, f"fullscreen_screenshot_{timestamp}.png")
    screenshot = pyautogui.screenshot()  # Fullscreen screenshot by default
    screenshot.save(screenshot_path)
    append_to_word(screenshot_path)  # Directly add to Word file
    print(f"Full-screen screenshot saved to {screenshot_path}")

# Crop GUI
def open_crop_gui(image_path):
    root = tk.Tk()
    root.title("Crop Screenshot")

    img = Image.open(image_path)
    tk_image = ImageTk.PhotoImage(img)

    canvas = tk.Canvas(root, width=img.width, height=img.height)
    canvas.pack()

    crop_coords = []
    current_rectangle = None

    def on_mouse_press(event):
        crop_coords.clear()  # Clear previous selections
        crop_coords.append((event.x, event.y))  # Initialize starting point for cropping

    def on_mouse_release(event):
        if len(crop_coords) == 1:  # Ensure there is a starting point
            crop_coords.append((event.x, event.y))  # Add the second coordinate when the mouse is released
            save_button.config(state="normal")  # Enable the Save button

    def on_mouse_drag(event):
        if len(crop_coords) == 1:  # Ensure there's a starting point
            crop_coords.append((event.x, event.y))  # Add the second coordinate during drag
        if len(crop_coords) == 2:  # Update the second coordinate during drag
            crop_coords[1] = (event.x, event.y)
            draw_rectangle(crop_coords)

    def save_selection(image_path):
        if len(crop_coords) == 2:  # Ensure we have both coordinates
            if crop_image(image_path, crop_coords):  # If crop_image returns True (successful)
                root.quit()  # Stop the mainloop after saving
                root.destroy()  # Close the Tkinter window

    def clear_selection():
        crop_coords.clear()
        canvas.delete("rectangle")  # Clear the rectangle
        canvas.bind("<ButtonPress-1>", on_mouse_press)
        canvas.bind("<ButtonRelease-1>", on_mouse_release)

    def draw_rectangle(crop_coords):
        # Only proceed if there are exactly 2 coordinates
        if len(crop_coords) == 2:
            (x1, y1), (x2, y2) = crop_coords
            canvas.delete("rectangle")  # Clear previous rectangle
            canvas.create_rectangle(x1, y1, x2, y2, outline="red", tags="rectangle")
    
    def capture_fullscreen_screenshot_gui():
        capture_fullscreen_screenshot()
        root.quit()
        root.destroy()

    # Create the buttons at the top-left corner
    save_button = tk.Button(root, text="Save", command=lambda: save_selection(image_path), state="disabled")
    save_button.place(x=10, y=10)

    clear_button = tk.Button(root, text="Clear Selection", command=clear_selection)
    clear_button.place(x=80, y=10)

    # Button to capture full-screen screenshot
    fullscreen_button = tk.Button(root, text="Capture Full-Screen", command=capture_fullscreen_screenshot_gui)
    fullscreen_button.place(x=210, y=10)

    canvas.create_image(0, 0, anchor=tk.NW, image=tk_image)

    # Bind mouse actions
    canvas.bind("<ButtonPress-1>", on_mouse_press)
    canvas.bind("<ButtonRelease-1>", on_mouse_release)
    canvas.bind("<B1-Motion>", on_mouse_drag)  # Update the rectangle during drag

    root.mainloop()

def crop_image(image_path, crop_coords):
    try:
        img = Image.open(image_path)
        (x1, y1), (x2, y2) = crop_coords

        # Ensure the crop coordinates are within the image bounds
        img_width, img_height = img.size
        x1 = max(0, min(x1, img_width))
        y1 = max(0, min(y1, img_height))
        x2 = max(0, min(x2, img_width))
        y2 = max(0, min(y2, img_height))

        # Ensure the coordinates are in the correct order (left < right, top < bottom)
        x1, x2 = sorted([x1, x2])
        y1, y2 = sorted([y1, y2])

        cropped_img = img.crop((x1, y1, x2, y2))
        save_path = f"{image_path}_cropped.png"  # You can customize this as needed
        cropped_img.save(save_path)
        append_to_word(save_path)

        print(f"Cropped image saved to {save_path}")
        return True
    except Exception as e:
        print(f"Error cropping image: {e}")
        return False

# Append to Word File
def append_to_word(image_path):
    if not os.path.exists(word_file_path):
        doc = Document()
    else:
        doc = Document(word_file_path)
    doc.add_picture(image_path, width=Inches(6))
    doc.save(word_file_path)
    print(f"Saved to {word_file_path}")

# Hotkey Listener
def hotkey_listener():
    def on_activate_screenshot():
        print("Hotkey pressed: Capturing screenshot...")
        capture_screenshot()

    def on_activate_fullscreen():
        print("Hotkey pressed: Capturing full-screen screenshot...")
        capture_fullscreen_screenshot()

    with keyboard.GlobalHotKeys({
        '<ctrl>+<shift>+x': on_activate_screenshot,
        '<ctrl>+<shift>+a': on_activate_fullscreen
    }) as h:
        h.join()

# Main
if __name__ == "__main__":
    threading.Thread(target=hotkey_listener, daemon=True).start()
    print("Application is running. Listening for hotkeys...")
    while True:
        pass  # Keep the script running
